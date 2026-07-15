"""
Plain-text extraction for admin-uploaded Bible Study source documents.

Admins have real devotional guides (WhatsApp broadcasts, Word docs, PDFs)
that they want turned into a structured multi-day StudyPlan. This module
only handles step one of that pipeline: turning whatever file they upload
into plain text. What happens to that text (the Gemini call, the JSON
shape it comes back as) lives in ai_assistant.py / bible.py — this file
has no opinion about AI at all, same spirit as ai_assistant.py staying
out of "where does the text get saved".
"""
import io
import logging

logger = logging.getLogger(__name__)

MAX_CHARS = 40_000  # keep prompts to Gemini reasonably sized

ALLOWED_EXTENSIONS = {"docx", "pdf", "txt", "md"}


class DocumentExtractError(Exception):
    """Raised when a file can't be read/converted into text."""


def _get_extension(filename: str) -> str | None:
    if not filename or "." not in filename:
        return None
    return filename.rsplit(".", 1)[1].lower()


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise DocumentExtractError(
            "Word document support isn't installed on the server (python-docx)."
        )
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"Failed to parse .docx: {e}")
        raise DocumentExtractError("Couldn't read that Word document — is it a valid .docx file?")

    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise DocumentExtractError(
            "PDF support isn't installed on the server (pypdf)."
        )
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = [page.extract_text() or "" for page in reader.pages]
    except Exception as e:
        logger.error(f"Failed to parse .pdf: {e}")
        raise DocumentExtractError("Couldn't read that PDF — is it a valid, non-scanned PDF?")
    text = "\n".join(parts)
    if not text.strip():
        raise DocumentExtractError(
            "That PDF doesn't seem to contain selectable text (it may be a scan/image)."
        )
    return text


def extract_text(file_storage) -> str:
    """
    file_storage: a werkzeug FileStorage from request.files['file'].
    Returns plain text, trimmed to MAX_CHARS.
    Raises DocumentExtractError with a user-facing message on failure.
    """
    filename = file_storage.filename or ""
    ext = _get_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise DocumentExtractError(
            f"Unsupported file type '.{ext or '?'}'. Upload a .docx, .pdf, .txt, or .md file."
        )

    file_bytes = file_storage.read()
    if not file_bytes:
        raise DocumentExtractError("Uploaded file is empty.")

    if ext == "docx":
        text = _extract_docx(file_bytes)
    elif ext == "pdf":
        text = _extract_pdf(file_bytes)
    else:  # txt / md
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="ignore")

    text = text.strip()
    if not text:
        raise DocumentExtractError("Couldn't find any readable text in that file.")

    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]

    return text
